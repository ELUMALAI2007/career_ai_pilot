"""
CareerPilot AI - Resume Parser Module
Extracts raw text, structured sections, contact metadata, and calculates completeness score from PDF and DOCX documents.
"""

import os
import re
from typing import Dict, Any, List, Tuple
from pypdf import PdfReader
from docx import Document


class ResumeParser:
    """Document text extractor and section parser for PDF and DOCX files."""

    ALLOWED_EXTENSIONS = {'pdf', 'docx'}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB

    @classmethod
    def validate_file(cls, file_path: str) -> Tuple[bool, str]:
        """Validates file extension, existence, size, and corruption."""
        if not os.path.exists(file_path):
            return False, "File does not exist on disk."

        file_size = os.path.getsize(file_path)
        if file_size > cls.MAX_FILE_SIZE:
            return False, f"File size ({file_size / (1024*1024):.1f} MB) exceeds maximum allowed limit of 50 MB."

        ext = file_path.rsplit('.', 1)[-1].lower() if '.' in file_path else ''
        if ext not in cls.ALLOWED_EXTENSIONS:
            return False, f"Unsupported file type '.{ext}'. Only PDF and DOCX files are supported."

        return True, "Valid"

    @classmethod
    def extract_text(cls, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extracts raw plain text and document metadata (page count, word count).
        Returns (text, metadata_dict).
        """
        is_valid, err_msg = cls.validate_file(file_path)
        if not is_valid:
            raise ValueError(err_msg)

        ext = file_path.rsplit('.', 1)[-1].lower()
        extracted_text = ""
        page_count = 1
        has_tables = False
        has_images = False

        if ext == 'pdf':
            try:
                reader = PdfReader(file_path)
                page_count = len(reader.pages)
                pages_text = []
                for idx, page in enumerate(reader.pages):
                    txt = page.extract_text() or ""
                    pages_text.append(txt)
                    if "/XObject" in str(page.get('/Resources', {})):
                        has_images = True
                extracted_text = "\n".join(pages_text)
            except Exception as e:
                raise ValueError(f"Failed to parse PDF document. File may be corrupted or encrypted: {str(e)}")

        elif ext == 'docx':
            try:
                doc = Document(file_path)
                paras = [p.text for p in doc.paragraphs if p.text.strip()]
                tables_text = []
                if len(doc.tables) > 0:
                    has_tables = True
                    for table in doc.tables:
                        for row in table.rows:
                            row_txt = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            if row_txt:
                                tables_text.append(row_txt)
                extracted_text = "\n".join(paras + tables_text)
                word_count = len(extracted_text.split())
                page_count = max(1, word_count // 350)  # Estimate 350 words per page
            except Exception as e:
                raise ValueError(f"Failed to parse DOCX document. File may be corrupted: {str(e)}")

        clean_text = re.sub(r'\r\n|\r', '\n', extracted_text).strip()
        if not clean_text or len(clean_text) < 20:
            raise ValueError("Extracted text is empty or unreadable. Ensure the document is not an image-only scan.")

        word_count = len(clean_text.split())
        metadata = {
            "page_count": page_count,
            "word_count": word_count,
            "has_tables": has_tables,
            "has_images": has_images,
            "file_ext": ext
        }
        return clean_text, metadata

    @classmethod
    def parse_contact_info(cls, text: str) -> Dict[str, Any]:
        """Extracts candidate contact details using pattern matching."""
        email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        phone_match = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
        linkedin_match = re.search(r'(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+/?', text, re.IGNORECASE)
        github_match = re.search(r'(https?://)?(www\.)?github\.com/[a-zA-Z0-9_-]+/?', text, re.IGNORECASE)
        portfolio_match = re.search(r'(https?://)?(www\.)?[a-zA-Z0-9_-]+\.(io|me|dev|com|org)/?', text, re.IGNORECASE)

        # Name extraction heuristic (first non-empty line)
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        candidate_name = lines[0] if lines else "Candidate"
        if len(candidate_name) > 50 or "@" in candidate_name or "Resume" in candidate_name:
            candidate_name = "Candidate Name"

        return {
            "name": candidate_name,
            "email": email_match.group(0) if email_match else None,
            "phone": phone_match.group(0) if phone_match else None,
            "linkedin": linkedin_match.group(0) if linkedin_match else None,
            "github": github_match.group(0) if github_match else None,
            "portfolio": portfolio_match.group(0) if portfolio_match else None,
        }

    @classmethod
    def parse_sections(cls, text: str) -> Dict[str, Any]:
        """Categorizes document text into standard resume sections."""
        lowered = text.lower()
        
        sections_found = {
            "summary": bool(re.search(r'(summary|profile|about me|objective)', lowered)),
            "education": bool(re.search(r'(education|academic|qualification|degree|university|college)', lowered)),
            "skills": bool(re.search(r'(skills|technical skills|technologies|competencies|expertise)', lowered)),
            "experience": bool(re.search(r'(experience|work history|employment|internship|work experience)', lowered)),
            "projects": bool(re.search(r'(projects|key projects|academic projects|personal projects)', lowered)),
            "certifications": bool(re.search(r'(certifications|certificates|licenses|courses)', lowered)),
            "achievements": bool(re.search(r'(achievements|honors|awards|accomplishments)', lowered)),
            "leadership": bool(re.search(r'(leadership|responsibility|extra-curricular|volunteering)', lowered)),
            "publications": bool(re.search(r'(publications|papers|patents)', lowered))
        }

        # Extract parsed skills using tech lexicon matching
        tech_lexicon = [
            "python", "java", "c++", "c#", "c", "javascript", "typescript", "html", "css", "sql", "postgresql",
            "mysql", "mongodb", "sqlite", "redis", "react", "next.js", "angular", "vue", "node.js", "express",
            "flask", "django", "fastapi", "spring boot", "docker", "kubernetes", "aws", "azure", "gcp",
            "git", "github", "linux", "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras",
            "power bi", "tableau", "excel", "spark", "hadoop", "airflow", "kafka", "rest api", "graphql",
            "machine learning", "deep learning", "nlp", "computer vision", "statistics", "data analysis"
        ]

        extracted_skills = []
        for skill in tech_lexicon:
            # Word boundary regex search
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, lowered):
                # Capitalize nicely
                extracted_skills.append(skill.title() if len(skill) > 3 else skill.upper())

        return {
            "sections_found": sections_found,
            "extracted_skills": sorted(list(set(extracted_skills))),
        }

    @classmethod
    def calculate_completeness_score(cls, contact: Dict[str, Any], sections: Dict[str, Any]) -> Tuple[float, Dict[str, Any]]:
        """
        Calculates Resume Completeness Score (0-100) with explicit breakdown.
        """
        breakdown = {}
        score = 0.0

        # Contact Info (20 points max)
        c_score = 0
        if contact.get("name") and contact["name"] != "Candidate Name": c_score += 4
        if contact.get("email"): c_score += 4
        if contact.get("phone"): c_score += 4
        if contact.get("linkedin"): c_score += 4
        if contact.get("github") or contact.get("portfolio"): c_score += 4
        breakdown["Contact & Links"] = {"score": c_score, "max": 20}
        score += c_score

        # Education (20 points max)
        sec = sections.get("sections_found", {})
        ed_score = 20 if sec.get("education") else 0
        breakdown["Education Section"] = {"score": ed_score, "max": 20}
        score += ed_score

        # Skills (20 points max)
        sk_count = len(sections.get("extracted_skills", []))
        sk_score = 20 if sk_count >= 8 else (sk_count * 2.5 if sk_count > 0 else 0)
        breakdown["Skills Section"] = {"score": sk_score, "max": 20}
        score += sk_score

        # Experience & Internships (20 points max)
        exp_score = 20 if sec.get("experience") else 0
        breakdown["Experience / Internships"] = {"score": exp_score, "max": 20}
        score += exp_score

        # Projects (15 points max)
        proj_score = 15 if sec.get("projects") else 0
        breakdown["Projects Section"] = {"score": proj_score, "max": 15}
        score += proj_score

        # Certifications & Achievements (5 points max)
        cert_score = 5 if (sec.get("certifications") or sec.get("achievements")) else 0
        breakdown["Certifications / Achievements"] = {"score": cert_score, "max": 5}
        score += cert_score

        return min(100.0, round(score, 1)), breakdown
